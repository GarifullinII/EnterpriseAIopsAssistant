from pathlib import Path
from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook
import xlrd
import subprocess
import tempfile
import shutil


def extract_text_from_txt_or_md(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def extract_text_from_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []

    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)

    return "\n".join(parts)


def extract_text_from_docx(path: Path) -> str:
    doc = Document(str(path))

    parts = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_text_from_xlsx(path: Path) -> str:
    wb = load_workbook(filename=str(path), read_only=True, data_only=True)
    parts = []

    for sheet in wb.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")

        for row in sheet.iter_rows(values_only=True):
            values = [str(cell) if cell is not None else "" for cell in row]
            if any(v.strip() for v in values):
                parts.append(" | ".join(values))

    return "\n".join(parts)


def extract_text_from_xls(path: Path) -> str:
    book = xlrd.open_workbook(str(path))
    parts = []

    for sheet in book.sheets():
        parts.append(f"[Sheet: {sheet.name}]")

        for row_idx in range(sheet.nrows):
            row = sheet.row_values(row_idx)
            values = [str(cell) if cell is not None else "" for cell in row]
            if any(v.strip() for v in values):
                parts.append(" | ".join(values))

    return "\n".join(parts)


def convert_doc_to_docx(path: Path) -> Path:
    # Требуется установленный LibreOffice / soffice
    with tempfile.TemporaryDirectory() as tmpdir:
        result = subprocess.run(
            [
                "soffice",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                tmpdir,
                str(path),
            ],
            check=True,
            capture_output=True,
            text=True,
        )

        converted = Path(tmpdir) / f"{path.stem}.docx"
        if not converted.exists():
            raise RuntimeError(f"Не удалось конвертировать .doc в .docx: {result.stdout} {result.stderr}")

        final_path = path.with_suffix(".converted.docx")
        final_path.write_bytes(converted.read_bytes())
        return final_path


def extract_text_from_doc(path: Path) -> str:
    converted_path = convert_doc_to_docx(path)
    try:
        return extract_text_from_docx(converted_path)
    finally:
        if converted_path.exists():
            converted_path.unlink()


def extract_text_from_file(file_path: str) -> str:
    path = Path(file_path)

    if not path.is_file():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix in {".txt", ".md"}:
        return extract_text_from_txt_or_md(path)

    if suffix == ".pdf":
        return extract_text_from_pdf(path)

    if suffix == ".docx":
        return extract_text_from_docx(path)

    if suffix == ".xlsx":
        return extract_text_from_xlsx(path)

    if suffix == ".xls":
        return extract_text_from_xls(path)

    if suffix == ".doc":
        if shutil.which("soffice") is None:
            raise ValueError(
                "Unsupported file type: .doc (LibreOffice/soffice is not installed)"
            )
        return extract_text_from_doc(path)

    raise ValueError(f"Unsupported file type: {suffix}")
